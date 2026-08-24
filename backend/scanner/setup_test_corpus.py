import os
import sys
import urllib.request
import docx

# Ensure parent directory of backend is on python path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))

CORPUS_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "data", "test_corpus"))

def setup_corpus():
    os.makedirs(CORPUS_DIR, exist_ok=True)
    print(f"Creating test corpus at: {CORPUS_DIR}")
    
    # 1. Create a text file
    txt_path = os.path.join(CORPUS_DIR, "doc1_python_intro.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(
            "Introduction to Python Programming.\n"
            "Python is a high-level, interpreted programming language known for its readability.\n"
            "It is widely used in data science, web development, and agentic AI systems.\n"
            "This file serves as a test document for Compass indexing."
        )
    print(f"Created text file: {txt_path}")
    
    # 2. Create a DOCX file
    docx_path = os.path.join(CORPUS_DIR, "doc2_search_strategies.docx")
    doc = docx.Document()
    doc.add_heading("Information Retrieval Strategies", level=1)
    doc.add_paragraph(
        "Modern search systems use multiple retrieval strategies to find relevant documents. "
        "The cheap path is keyword search, which matches exact tokens using TF-IDF or BM25 models. "
        "The expensive path is semantic search, which represents text as high-dimensional vectors and "
        "compares them using cosine similarity. Hybrid search merges results from both paths."
    )
    doc.save(docx_path)
    print(f"Created Word file: {docx_path}")
    
    # 3. Download a sample PDF file
    # We download a small, standard dummy PDF from a public W3C test resource
    pdf_path = os.path.join(CORPUS_DIR, "doc3_dummy_test.pdf")
    pdf_url = "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf"
    try:
        print(f"Downloading sample PDF from: {pdf_url}")
        urllib.request.urlretrieve(pdf_url, pdf_path)
        print(f"Downloaded PDF file: {pdf_path}")
    except Exception as e:
        print(f"Warning: Failed to download sample PDF: {e}")
        print("Creating a fallback txt-based PDF (which will fail extraction but scanner should handle it gracefully)")
        with open(pdf_path, "w") as f:
            f.write("%PDF-1.4 (dummy fallback)")
            
    print("Test corpus setup completed successfully.")

if __name__ == "__main__":
    setup_corpus()
